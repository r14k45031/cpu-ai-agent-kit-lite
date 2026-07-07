# -*- coding: utf-8 -*-
"""Word 就地 AI 改寫：逐段改寫後寫回原版面——樣式、表格、圖片全部保留。

用法：edit_docx_inplace.py <docx 路徑>（執行時會詢問改寫指示）
輸出：<原檔名>_就地改寫.docx（原檔不動）

保留：段落樣式、標題階層、表格結構、圖片、頁首頁尾。
限制：被改寫段落內的「字元級」格式（如句中粗體）會統一成該段第一種格式。
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
import ai_common as ai

DEFAULT = "潤稿：使語句通順、專業、簡潔，修正錯字與贅詞，保留原意"


def iter_paragraphs(doc):
    """依序取出內文與表格內的所有段落。"""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def should_rewrite(text):
    """太短的段落（標題、表格數據格）不送 AI，避免數字被改壞。"""
    if len(text) < 15:
        return False
    digits = sum(c.isdigit() or c in ",.%$￥：:/-" for c in text)
    return digits / len(text) < 0.3


def rewrite_paragraph(p, instr, model):
    text = p.text.strip()
    if not p.runs:
        return False
    new = ai.chat(
        "請依照指示改寫以下段落。要求：一律使用繁體中文（台灣用語）；"
        "使用正常的標點符號；段落中的數字、金額、日期必須原封不動；"
        f"只輸出改寫後的段落文字，不要任何說明。\n指示：{instr}\n\n{text}",
        model=model,
    )
    new = new.strip().replace("\n", " ")
    if not new:
        return False
    p.runs[0].text = new
    for r in p.runs[1:]:
        r.text = ""
    return True


def main():
    if len(sys.argv) < 2:
        raise SystemExit("用法：把 .docx 拖到「拖入Word就地AI改寫.bat」上")
    f = Path(sys.argv[1])
    if f.suffix.lower() != ".docx":
        raise SystemExit("請拖入 .docx 檔案")
    import docx

    try:
        instr = input(f"要怎麼改寫？（直接按 Enter ＝ {DEFAULT}）\n> ").strip()
    except EOFError:
        instr = ""
    instr = instr or DEFAULT

    model = ai.pick_model()
    doc = docx.Document(str(f))
    paras = [p for p in iter_paragraphs(doc) if should_rewrite(p.text.strip())]
    print(f"讀取：{f.name}，共 {len(paras)} 個段落要改寫（模型：{model}）")
    done = 0
    for i, p in enumerate(paras, 1):
        print(f"  [{i}/{len(paras)}] {p.text.strip()[:20]}...")
        try:
            if rewrite_paragraph(p, instr, model):
                done += 1
        except Exception as e:
            print(f"    此段失敗，保留原文：{e}")
    out = f.with_name(f.stem + "_就地改寫.docx")
    doc.save(str(out))
    print(f"完成：改寫 {done}/{len(paras)} 段，版面與表格圖片皆保留。已輸出：{out.name}")


if __name__ == "__main__":
    main()
